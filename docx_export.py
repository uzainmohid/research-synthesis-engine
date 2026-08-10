"""
DOCX Exporter (Step 6)
==========================
Renders the assembled report (executive summary, findings, conflicts &
gaps, references) into a clean, professional .docx file using
python-docx.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_docx(question: str, report: dict, claims: list, path: str) -> str:
    doc = Document()

    # Title
    title = doc.add_heading("Research Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(question)
    run.italic = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(report.get("executive_summary", ""))

    # Findings
    doc.add_heading("Findings", level=1)
    for i, finding in enumerate(report.get("findings", []), 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(finding)

    # Conflicts & Gaps
    doc.add_heading("Conflicts & Gaps", level=1)
    conflicts = report.get("conflicts_and_gaps", [])
    if conflicts:
        for c in conflicts:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(c)
    else:
        doc.add_paragraph("No conflicts identified across sources for this question.")

    # References
    doc.add_heading("References", level=1)
    for num, source_name, url in report.get("references", []):
        p = doc.add_paragraph()
        p.add_run(f"[{num}] ").bold = True
        p.add_run(f"{source_name}: ")
        p.add_run(url if url else "(no URL available)")

    # Appendix: raw claim data with quality scores (transparency)
    doc.add_page_break()
    doc.add_heading("Appendix: All Extracted Claims", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Claim"
    hdr[1].text = "Source"
    hdr[2].text = "Quality Score"
    hdr[3].text = "Retrieved At"
    for claim in claims:
        row = table.add_row().cells
        row[0].text = claim.claim_text
        row[1].text = claim.source_name
        row[2].text = str(claim.quality_score)
        row[3].text = claim.retrieved_at

    doc.save(path)
    return path
